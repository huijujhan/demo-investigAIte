import streamlit as st
import boto3
import json
import os
import re
import uuid
import pandas as pd
from io import BytesIO
from datetime import datetime
from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt

load_dotenv()

# ── Page config ────────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="InvestigAIte",
    page_icon="🔍",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ── Custom CSS ─────────────────────────────────────────────────────────────────
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=IBM+Plex+Mono:wght@400;500&family=IBM+Plex+Sans:wght@300;400;600&display=swap');

:root {
    --bg:        #ffffff;
    --surface:   #f8f9fa;
    --border:    #e2e6ea;
    --accent:    #1a6ef5;
    --accent2:   #0d9e6e;
    --text:      #1a1d23;
    --muted:     #6c757d;
}

html, body, [class*="css"] {
    font-family: 'IBM Plex Sans', sans-serif;
    background-color: var(--bg);
    color: var(--text);
}

#MainMenu, footer, header { visibility: hidden; }
.block-container { padding: 1.5rem 2rem 6rem; max-width: 900px; margin: auto; }

/* Logo area */
.logo-area{
    display: flex; justify-content: center; align-items: center;
    padding: 24px 0 8px;
    border-bottom: 2px solid var(--border);
    margin-bottom: 8px;
}
.logo-placeholder{
    font-size: 1.6rem; font-weight: 700; color: var(--accent);
    font-family: 'IBM Plex Sans', sans-serif; letter-spacing: -0.5px;
}
.tagline {
    text-align: center; font-size: 0.78rem; color: var(--muted);
    font-family: 'IBM Plex Mono', monospace; margin-bottom: 24px;
}
            
/* Messages */
.badge {
    display: inline-block; font-size: 0.65rem; font-family: 'IBM Plex Mono', monospace;
    padding: 2px 7px; border-radius: 4px; margin-bottom: 6px; font-weight: 500;
}
.badge.arcos { background: rgba(79,255,176,0.12); color: var(--accent); border: 1px solid rgba(79,255,176,0.25); }
.badge.cfr   { background: rgba(10,240,255,0.12); color: var(--accent2); border: 1px solid rgba(10,240,255,0.25); }
.badge.dea6  { background: rgba(255,180,80,0.12);  color: #c97a00;       border: 1px solid rgba(255,180,80,0.25); }

/* Sidebar */
section[data-testid="stSidebar"] {
    background: var(--surface) !important;
    border-right: 1px solid var(--border) !important;
}
section[data-testid="stSidebar"] * { color: var(--text) !important; }

.sidebar-label {
    font-family: 'IBM Plex Mono', monospace;
    font-size: 0.65rem; color: var(--muted) !important;
    text-transform: uppercase; letter-spacing: 0.08em;
    margin-bottom: 4px;
}

/* Input */
.stChatInput > div { background: var(--surface) !important; border: 1px solid var(--border) !important; border-radius: 12px !important; }
.stChatInput textarea { color: var(--text) !important; font-family: 'IBM Plex Sans', sans-serif !important; }

/* Tables */
.stDataFrame { border: 1px solid var(--border) !important; border-radius: 8px !important; }

/* Ensure wide markdown tables scroll inside message content */
div[data-testid="stMarkdownContainer"] table {
    display: block;
    overflow-x: auto;
    max-width: 100%;
    white-space: nowrap;
}

/* Expander */
.streamlit-expanderHeader { background: var(--surface) !important; border: 1px solid var(--border) !important; border-radius: 8px !important; }
</style>
""", unsafe_allow_html=True)

# ── AWS config ─────────────────────────────────────────────────────────────────
# CFR is intentionally NOT a routing target. CFR regulations are surfaced as
# inline citations inside the ARCOS and DEA-6 lambda responses.
AWS_REGION          = os.getenv("AWS_REGION", "us-east-1")
ARCOS_LAMBDA_NAME   = os.getenv("ARCOS_LAMBDA_NAME", "arcos-nl-sql-handler")
DEA6_LAMBDA_NAME    = os.getenv("DEA6_LAMBDA_NAME",  "dea6-bedrock-kb-handler")

# ── Session state ──────────────────────────────────────────────────────────────
if "messages" not in st.session_state:
    st.session_state.messages = []
if "report_items" not in st.session_state:
    st.session_state.report_items = []
if "final_report" not in st.session_state:
    st.session_state.final_report = ""
if "report_template" not in st.session_state:
    st.session_state.report_template = ""  # set after DEFAULT_REPORT_TEMPLATE is defined

# ── History helpers ────────────────────────────────────────────────────────────
def build_history_payload(n: int = 6) -> list:
    """Return last n messages trimmed so first message is always a user turn."""
    history = [
        {"role": m["role"], "content": m["content"]}
        for m in st.session_state.messages[-n:]
        if m["role"] in ("user", "assistant")
    ]
    while history and history[0]["role"] != "user":
        history.pop(0)
    return history

# ── Lambda caller ──────────────────────────────────────────────────────────────
def invoke_lambda(function_name: str, payload: dict) -> dict:
    client = boto3.client("lambda", region_name=AWS_REGION)
    response = client.invoke(
        FunctionName=function_name,
        InvocationType="RequestResponse",
        Payload=json.dumps(payload).encode(),
    )
    return json.loads(response["Payload"].read())

# ── Intent classifier ──────────────────────────────────────────────────────────
INTENT_MODEL_ID = os.getenv("INTENT_MODEL_ID", "us.anthropic.claude-sonnet-4-6")

INTENT_SYSTEM_PROMPT = """You route user questions to one of two data sources.
Reply with exactly one lowercase token and nothing else: arcos or dea6.

- arcos: ARCOS controlled-substance transaction records (shipments, quantities,
  buyers, sellers, drug codes, dates, aggregations, SQL-style analytics, and
  local-area purchasing trends).
- dea6:  DEA-6 investigative reports (prior investigations and follow-up
  investigations of specific registrants, practitioners, pharmacies,
  distributors; narrative case findings; investigator notes; physical-
  security history; prior violations).

CFR (Code of Federal Regulations) is NOT a routing target. Regulatory
citations (e.g. 21 CFR 1301.71) are surfaced inside the ARCOS and DEA-6
responses as references, not as a separate query.

If the question asks purely about a regulation in the abstract, prefer
dea6 — its handler will cite the relevant CFR sections.

If ambiguous, prefer the source most consistent with the conversation so far."""

VALID_INTENTS = {"arcos", "dea6"}

@st.cache_resource
def _bedrock_client():
    return boto3.client("bedrock-runtime", region_name=AWS_REGION)

def classify_intent(question: str, history: list) -> str:
    """Use Bedrock to pick the best data source for this question."""
    messages = []
    for turn in history[-4:]:
        if turn["role"] in ("user", "assistant"):
            messages.append({"role": turn["role"], "content": turn["content"]})
    while messages and messages[0]["role"] != "user":
        messages.pop(0)
    messages.append({"role": "user", "content": question})

    try:
        resp = _bedrock_client().invoke_model(
            modelId=INTENT_MODEL_ID,
            body=json.dumps({
                "anthropic_version": "bedrock-2023-05-31",
                "max_tokens": 8,
                "temperature": 0,
                "system":   INTENT_SYSTEM_PROMPT,
                "messages": messages,
            }),
        )
        text = json.loads(resp["body"].read())["content"][0]["text"].strip().lower()
        for token in VALID_INTENTS:
            if token in text:
                return token
    except Exception:
        pass
    return "arcos"

# ── SQL → plain English ────────────────────────────────────────────────────────
EXPLAIN_MODEL_ID = os.getenv("EXPLAIN_MODEL_ID", INTENT_MODEL_ID)

EXPLAIN_SYSTEM_PROMPT = """You translate Athena SQL queries against the DEA ARCOS
opioid shipment database into a short, plain-English explanation of the query's
logic for a non-technical reader (journalist, policy analyst).

Rules:
- 2 to 5 short sentences or bullet points.
- Describe WHAT the query is asking for, not HOW SQL works.
- Mention filters (transaction type, drug, dates, location, buyer type),
  what is being measured (count, total dosage units, weight), how results are
  grouped, the sort order, and the row limit if any.
- Do NOT mention table or column names, SQL keywords, or syntax.
- Do NOT add a preamble like "This query...". Start directly with the logic."""

@st.cache_data(show_spinner=False)
def explain_sql(sql: str, question: str = "") -> str:
    """Generate a plain-English explanation of an Athena SQL query."""
    if not sql:
        return ""
    try:
        resp = _bedrock_client().invoke_model(
            modelId=EXPLAIN_MODEL_ID,
            body=json.dumps({
                "anthropic_version": "bedrock-2023-05-31",
                "max_tokens": 400,
                "temperature": 0,
                "system":   EXPLAIN_SYSTEM_PROMPT,
                "messages": [{
                    "role": "user",
                    "content": (
                        f"User question: {question}\n\n"
                        f"SQL:\n{sql}\n\n"
                        "Explain the query logic in plain English."
                    ),
                }],
            }),
        )
        return json.loads(resp["body"].read())["content"][0]["text"].strip()
    except Exception as e:
        return f"_Could not generate explanation: {e}_"

# ── Report compositor ──────────────────────────────────────────────────────────
REPORT_MODEL_ID = os.getenv("REPORT_MODEL_ID", INTENT_MODEL_ID)

REPORT_SYSTEM_PROMPT = """You are an investigative writing assistant. You
compose a structured markdown report from a set of saved findings produced
during an investigative chat session over ARCOS shipment data, CFR regulations,
DEA-6 investigative reports, and user-uploaded supporting documents.

Hard rules:
- Output valid GitHub-flavored markdown only. No preamble, no explanation,
  no surrounding code fences.
- If the user supplies a REQUIRED REPORT STRUCTURE and STARTING SKELETON,
  copy every heading from the skeleton VERBATIM (same text, same level,
  same order) and write the body of each section beneath it. Do not add,
  remove, rename, reorder, or change the level of any heading. Do not
  introduce new top-level headings.
- Guidance lines (the "Guidance:" notes next to each heading) are
  instructions for what to look for in the findings. NEVER copy guidance
  text into the final report.
- If a section has no supporting evidence in the findings, write
  "_No information available._" under that heading and move on.
- Preserve tables verbatim. Preserve key numbers, names, DEA numbers, and
  dates exactly as written.
- When a finding is sourced from an uploaded document, cite its file name
  inline (e.g. "(from `report.docx`)").
- Do not invent facts."""

DEFAULT_REPORT_TEMPLATE = """# Report of Investigation

## Background

### Registrant Information
> Registrant name, DEA registration number, business address, registered
> activity/schedules, registration status and expiration, point of contact,
> and any related corporate entities.

### Nature of Business
> Type of registrant (retail pharmacy, chain pharmacy, hospital/clinic,
> practitioner, distributor, manufacturer), patient/customer profile,
> typical controlled-substance handling, hours of operation, and any
> notable business relationships (suppliers, downstream buyers).

### Basis of Investigation
> Why this investigation was initiated: complaints, tips, ARCOS anomalies,
> prior compliance history, referral source, scheduling drugs of interest,
> and the time period under review.

## Investigation Narrative

### Pre-Inspection Activity
> Records reviewed before the site visit: ARCOS shipment trends, prior
> inspection findings, state licensing data, prescriber/patient overlap
> analyses, applicable CFR citations, and the investigative plan.

### On-site Inspection
> Date and personnel present, scope of the inspection, interviews
> conducted, sampling methodology, records examined (invoices, 222 forms,
> dispensing logs), and a chronological narrative of activities performed
> on site.

### Physical Security Review
> Assessment of vaults/safes, alarm systems, access controls, storage of
> Schedule II–V substances, recordkeeping/destruction practices, and
> compliance with 21 CFR Part 1301 security requirements.

## Violations / Deficiencies
> Itemized list of each violation or deficiency observed, with the
> specific CFR citation, supporting evidence (quantities, dates, missing
> records), severity, and any corrective action discussed with the
> registrant.
"""

if not st.session_state.report_template:
    st.session_state.report_template = DEFAULT_REPORT_TEMPLATE

INTENT_LABELS = {
    "arcos":  "ARCOS Data",
    "cfr":    "CFR Regulations",
    "dea6":   "DEA-6 Reports",
    "upload": "Uploaded Document",
}

def _format_staged_items(items: list) -> str:
    parts = []
    for i, item in enumerate(items, start=1):
        label = INTENT_LABELS.get(item.get("intent", ""), "Finding")
        header = f"### Finding {i} — {label}"
        if item.get("source_name"):
            header += f" (from `{item['source_name']}`)"
        parts.append(f"{header}\n{item['content']}")
        if item.get("sql_explanation"):
            parts.append(f"_Query logic:_ {item['sql_explanation']}")
        if item.get("sql"):
            parts.append(f"```sql\n{item['sql']}\n```")
    return "\n\n".join(parts)

def _parse_template(template: str) -> tuple[str, list[dict]]:
    """Split a markdown template into (skeleton, sections-with-guidance).

    Skeleton preserves every heading line verbatim. Sections is a list of
    {heading, guidance} dicts where guidance is the concatenation of any
    `>` blockquote lines that follow the heading.
    """
    lines = template.splitlines()
    skeleton_lines: list[str] = []
    sections: list[dict] = []
    current: dict | None = None
    for line in lines:
        stripped = line.strip()
        if stripped.startswith("#"):
            skeleton_lines.append(stripped)
            current = {"heading": stripped, "guidance": []}
            sections.append(current)
        elif stripped.startswith(">") and current is not None:
            current["guidance"].append(stripped.lstrip("> ").rstrip())
        # ignore other lines (blank, etc.) for skeleton purposes
    for s in sections:
        s["guidance"] = " ".join(s["guidance"]).strip()
    skeleton = "\n\n".join(skeleton_lines)
    return skeleton, sections

def compose_report(items: list, title_hint: str = "", template: str = "") -> str:
    """Use Bedrock to synthesize a markdown report from staged findings."""
    if not items:
        return ""

    sections_block = ""
    skeleton = ""
    if template.strip():
        skeleton, parsed = _parse_template(template)
        guidance_lines = []
        for s in parsed:
            line = f"- {s['heading']}"
            if s["guidance"]:
                line += f"\n  Guidance: {s['guidance']}"
            guidance_lines.append(line)
        sections_block = (
            "REQUIRED REPORT STRUCTURE (use these exact headings, in this exact "
            "order, and no others):\n\n"
            + "\n".join(guidance_lines)
            + "\n\nSTARTING SKELETON (fill in the body under each heading, do "
              "NOT change, reorder, add, or remove any heading):\n\n"
            + skeleton
        )

    parts = []
    if title_hint:
        parts.append(f"Suggested title: {title_hint}")
    if sections_block:
        parts.append(sections_block)
    parts.append("SAVED FINDINGS (use these as the only source of facts):\n\n"
                 + _format_staged_items(items))
    parts.append(
        "Compose the final markdown report now. If the structure above is "
        "provided, your output MUST start with the first heading from the "
        "skeleton and contain every heading from the skeleton in order, with "
        "the body of each section drawn from the saved findings. If a section "
        "has no supporting evidence, write \"_No information available._\" "
        "under it. Do not output the guidance lines."
    )
    user_msg = "\n\n".join(parts)
    try:
        resp = _bedrock_client().invoke_model(
            modelId=REPORT_MODEL_ID,
            body=json.dumps({
                "anthropic_version": "bedrock-2023-05-31",
                "max_tokens": 4000,
                "temperature": 0.2,
                "system":   REPORT_SYSTEM_PROMPT,
                "messages": [{"role": "user", "content": user_msg}],
            }),
        )
        return json.loads(resp["body"].read())["content"][0]["text"].strip()
    except Exception as e:
        return f"_Report generation failed: {e}_"

# ── Markdown → Word converter ──────────────────────────────────────────────────
INLINE_PATTERN = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)")

def _add_inline_runs(paragraph, text: str) -> None:
    """Render simple inline markdown (bold, italic, code) into a paragraph."""
    for chunk in INLINE_PATTERN.split(text):
        if not chunk:
            continue
        if chunk.startswith("**") and chunk.endswith("**"):
            run = paragraph.add_run(chunk[2:-2]); run.bold = True
        elif chunk.startswith("*") and chunk.endswith("*"):
            run = paragraph.add_run(chunk[1:-1]); run.italic = True
        elif chunk.startswith("`") and chunk.endswith("`"):
            run = paragraph.add_run(chunk[1:-1]); run.font.name = "Consolas"
        else:
            paragraph.add_run(chunk)

def markdown_to_docx(md: str) -> bytes:
    """Convert a subset of markdown (headings, lists, tables, code, paragraphs) to .docx."""
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    lines = md.splitlines()
    i = 0
    in_code = False
    code_buf: list = []
    while i < len(lines):
        line = lines[i]

        # Fenced code blocks
        if line.strip().startswith("```"):
            if in_code:
                p = doc.add_paragraph()
                run = p.add_run("\n".join(code_buf))
                run.font.name = "Consolas"; run.font.size = Pt(10)
                code_buf = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_buf.append(line)
            i += 1
            continue

        stripped = line.strip()

        # Headings
        if stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            level = min(max(level, 1), 4)
            doc.add_heading(stripped[level:].strip(), level=level)
            i += 1
            continue

        # Tables (consecutive | ... | lines, with separator row)
        if stripped.startswith("|") and stripped.endswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            rows = [
                [c.strip() for c in row.strip("|").split("|")]
                for row in table_lines
                if not re.match(r"^\|?\s*:?-+:?\s*(\|\s*:?-+:?\s*)+\|?$", row)
            ]
            if rows:
                table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                table.style = "Light Grid Accent 1"
                for r, row in enumerate(rows):
                    for c, cell in enumerate(row):
                        if c < len(table.rows[r].cells):
                            table.rows[r].cells[c].text = cell
            continue

        # Bullet lists
        if stripped.startswith(("- ", "* ", "+ ")):
            p = doc.add_paragraph(style="List Bullet")
            _add_inline_runs(p, stripped[2:])
            i += 1
            continue

        # Numbered lists
        if re.match(r"^\d+\.\s", stripped):
            p = doc.add_paragraph(style="List Number")
            _add_inline_runs(p, re.sub(r"^\d+\.\s", "", stripped))
            i += 1
            continue

        # Blank line
        if not stripped:
            doc.add_paragraph("")
            i += 1
            continue

        # Default paragraph
        p = doc.add_paragraph()
        _add_inline_runs(p, line)
        i += 1

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()

def docx_to_text(file_bytes: bytes) -> str:
    """Extract paragraphs and tables from a .docx as plain markdown-ish text."""
    doc = Document(BytesIO(file_bytes))
    out: list = []
    for block in doc.element.body.iterchildren():
        tag = block.tag.split("}")[-1]
        if tag == "p":
            text = "".join(node.text or "" for node in block.iter() if node.tag.endswith("}t")).strip()
            if text:
                out.append(text)
        elif tag == "tbl":
            rows = []
            for tr in block.iter("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tr"):
                cells = [
                    " ".join((node.text or "") for node in tc.iter() if node.tag.endswith("}t")).strip()
                    for tc in tr.iter("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tc")
                ]
                rows.append("| " + " | ".join(cells) + " |")
            if rows:
                if len(rows) >= 1:
                    sep = "| " + " | ".join(["---"] * rows[0].count("|") if rows[0].count("|") > 0 else ["---"]) + " |"
                    out.append(rows[0])
                    out.append(sep)
                    out.extend(rows[1:])
    return "\n\n".join(out).strip()

# ── Main query handler ─────────────────────────────────────────────────────────
def handle_query(question: str) -> dict:
    history = build_history_payload(n=6)
    intent  = classify_intent(question, history)

    lambda_map = {
        "arcos": ARCOS_LAMBDA_NAME,
        "dea6":  DEA6_LAMBDA_NAME,
    }

    try:
        result = invoke_lambda(lambda_map[intent], {"question": question, "history": history})
        response = {
            "answer": result.get("answer", "No answer returned."),
            "intent": intent,
        }
        if intent == "arcos":
            response["sql"] = result.get("sql")
            response["sql_explanation"] = result.get("sql_explanation")
            response["columns"] = result.get("columns") or []
            response["rows"]    = result.get("rows") or []
        return response
    except Exception as e:
        return {"answer": f"⚠️ Error calling Lambda: {str(e)}", "intent": intent}


def render_query_results(columns: list, rows: list) -> None:
    """Render Athena results as a dataframe, plus a bar chart when appropriate."""
    if not columns or not rows:
        return

    df = pd.DataFrame(rows, columns=columns)
    # Coerce columns that parse as numeric (more than half of values)
    for col in df.columns:
        coerced = pd.to_numeric(df[col], errors="coerce")
        if coerced.notna().sum() >= max(1, len(df) // 2 + 1):
            df[col] = coerced

    st.markdown("**Query results**")
    st.dataframe(df, use_container_width=True, hide_index=True)

    # Chart when there's exactly one non-numeric label column and >=1 numeric column,
    # and the result set isn't huge.
    if len(df) > 30:
        return
    numeric_cols = [c for c in df.columns if pd.api.types.is_numeric_dtype(df[c])]
    label_cols   = [c for c in df.columns if c not in numeric_cols]
    if len(label_cols) == 1 and len(numeric_cols) >= 1:
        chart_df = df.set_index(label_cols[0])[numeric_cols]
        st.bar_chart(chart_df, use_container_width=True)

# ── Sidebar ────────────────────────────────────────────────────────────────────
with st.sidebar:
    st.markdown("### ⚙️ Configuration")
    st.markdown("---")

    st.markdown('<p class="sidebar-label">AWS Region</p>', unsafe_allow_html=True)
    st.code(AWS_REGION, language=None)

    st.markdown('<p class="sidebar-label">ARCOS Lambda</p>', unsafe_allow_html=True)
    st.code(ARCOS_LAMBDA_NAME, language=None)

    st.markdown('<p class="sidebar-label">DEA-6 Lambda</p>', unsafe_allow_html=True)
    st.code(DEA6_LAMBDA_NAME, language=None)

    st.markdown("---")
    st.markdown(f'<p class="sidebar-label">Messages in session: {len(st.session_state.messages)}</p>', unsafe_allow_html=True)
    st.markdown(f'<p class="sidebar-label">Saved to report: {len(st.session_state.report_items)}</p>', unsafe_allow_html=True)

    if st.button("🗑️ Clear history", use_container_width=True):
        st.session_state.messages = []
        st.rerun()

# ── Logo ─────────────────────────────────────────────────────────────────────
LOGO_PATH = "InvestigAIte.png"
if os.path.exists(LOGO_PATH):
    st.image(LOGO_PATH, use_container_width=True)
else:
    st.markdown('<div class="logo-area"><span class="logo-placeholder">🔍 InvestigAIte</span></div>)',unsafe_allow_html=True)

st.markdown('<p class="tagline">ARCOS sales data - DEA-6 investigative reports</p>', unsafe_allow_html=True)

# ── Render chat history ────────────────────────────────────────────────────────
saved_ids = {item["id"] for item in st.session_state.report_items}

for msg in st.session_state.messages:
    role    = msg["role"]
    content = msg["content"]
    intent  = msg.get("intent", "")
    sql     = msg.get("sql")
    sql_explanation = msg.get("sql_explanation")
    columns = msg.get("columns") or []
    rows    = msg.get("rows") or []
    ts      = msg.get("ts", "")
    msg_id  = msg.setdefault("id", str(uuid.uuid4()))

    if role == "user":
        with st.chat_message("user", avatar="🕵️"):
            st.markdown(content)
            if ts:
                st.caption(ts)
    else:
        badge_map = {
            "arcos": ("arcos", "📊 ARCOS Data"),
            "dea6":  ("dea6",  "🔍 DEA-6 Reports"),
            # legacy: prior sessions may still have "cfr" tagged messages
            "cfr":   ("cfr",   "📋 CFR Regulations"),
        }
        badge_class, badge_label = badge_map.get(intent, ("arcos", "📊 ARCOS Data"))

        with st.chat_message("assistant", avatar="🤖"):
            st.markdown(
                f'<span class="badge {badge_class}">{badge_label}</span>',
                unsafe_allow_html=True,
            )
            if columns and rows:
                render_query_results(columns, rows)
            st.markdown(content)
            if ts:
                st.caption(ts)
            if sql:
                with st.expander("🔎 Query details"):
                    logic_tab, sql_tab = st.tabs(["📝 Query logic", "🧑‍💻 SQL"])
                    with logic_tab:
                        explanation = sql_explanation or explain_sql(sql, content)
                        if explanation:
                            st.markdown(explanation)
                        else:
                            st.caption("No plain-English explanation available.")
                    with sql_tab:
                        st.code(sql, language="sql")

            # Save to report staging area
            already_saved = msg_id in saved_ids
            save_label = "✅ Saved to report" if already_saved else "📌 Save to report"
            if st.button(
                save_label,
                key=f"save_{msg_id}",
                disabled=already_saved,
                help="Add this response to the report staging area.",
            ):
                st.session_state.report_items.append({
                    "id":      msg_id,
                    "intent":  intent,
                    "content": content,
                    "sql":     sql,
                    "sql_explanation": sql_explanation,
                    "ts":      ts,
                })
                st.rerun()

# ── Clear chat ─────────────────────────────────────────────────────────────────
if st.session_state.messages:
    _, clear_col = st.columns([5, 1])
    with clear_col:
        if st.button("🗑️ Clear chat", use_container_width=True, key="clear_chat_main"):
            st.session_state.messages = []
            st.rerun()

# ── Report staging area ─────────────────────────────────────────────────────────
staged = st.session_state.report_items
with st.expander(f"📄 Report staging area ({len(staged)})", expanded=False):
    findings_tab, template_tab = st.tabs(["📌 Findings", "🧱 Report template"])

    with findings_tab:
        # Uploader: add Word docs as additional findings
        uploads = st.file_uploader(
            "Upload Word documents to include in the report",
            type=["docx"],
            accept_multiple_files=True,
            key="report_uploader",
        )
        if uploads:
            existing_uploads = {
                item.get("source_name") for item in staged if item.get("intent") == "upload"
            }
            added = 0
            for f in uploads:
                if f.name in existing_uploads:
                    continue
                try:
                    text = docx_to_text(f.getvalue())
                except Exception as e:
                    st.warning(f"Could not read {f.name}: {e}")
                    continue
                if not text:
                    st.warning(f"{f.name} contained no readable text.")
                    continue
                staged.append({
                    "id":      str(uuid.uuid4()),
                    "intent":  "upload",
                    "content": text,
                    "source_name": f.name,
                    "sql":     None,
                    "sql_explanation": None,
                    "ts":      datetime.now().strftime("%H:%M"),
                })
                added += 1
            if added:
                st.success(f"Added {added} uploaded document(s) to the staging area.")
                st.rerun()

        if not staged:
            st.caption("No findings saved yet. Use the 📌 Save to report button under any AI response, or upload a Word document above.")
        else:
            for i, item in enumerate(staged):
                label = INTENT_LABELS.get(item.get("intent", ""), "Finding")
                preview_source = item.get("source_name") or item["content"].strip().splitlines()[0]
                preview = preview_source[:120]
                row_main, row_up, row_down, row_del = st.columns([8, 1, 1, 1])
                with row_main:
                    st.markdown(f"**{i+1}. {label}** — {preview}")
                with row_up:
                    if st.button("⬆️", key=f"up_{item['id']}", disabled=(i == 0), help="Move up"):
                        staged[i-1], staged[i] = staged[i], staged[i-1]
                        st.rerun()
                with row_down:
                    if st.button("⬇️", key=f"down_{item['id']}", disabled=(i == len(staged)-1), help="Move down"):
                        staged[i+1], staged[i] = staged[i], staged[i+1]
                        st.rerun()
                with row_del:
                    if st.button("❌", key=f"del_{item['id']}", help="Remove"):
                        staged.pop(i)
                        st.rerun()

    with template_tab:
        st.caption(
            "Edit the markdown skeleton below. The generator will fill each section "
            "using the saved findings. Use `#` for the title, `##` for sections, "
            "`###` for subsections. Lines starting with `>` are guidance for the "
            "generator and won't appear in the final report."
        )
        if st.button("↩️ Reset to default template", key="reset_template"):
            st.session_state["report_template"] = DEFAULT_REPORT_TEMPLATE
            st.rerun()
        st.text_area(
            "Report layout (markdown)",
            height=320,
            key="report_template",
        )

    st.markdown("---")
    title_hint = st.text_input("Report title (optional)", key="report_title")
    gen_col, clear_col = st.columns([1, 1])
    with gen_col:
        if st.button(
            "📝 Generate report",
            use_container_width=True,
            type="primary",
            disabled=not staged,
        ):
            with st.spinner("Composing report..."):
                st.session_state.final_report = compose_report(
                    staged,
                    title_hint,
                    st.session_state.report_template,
                )
    with clear_col:
        if st.button("🗑️ Clear", use_container_width=True, disabled=not staged):
            st.session_state.report_items = []
            st.session_state.final_report = ""
            st.rerun()

    if st.session_state.final_report:
        st.markdown("### Final report")
        st.markdown(st.session_state.final_report)
        st.download_button(
            "⬇️ Download as Word document",
            data=markdown_to_docx(st.session_state.final_report),
            file_name=f"investigaite_report_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
        )

# ── Chat input ─────────────────────────────────────────────────────────────────
if prompt := st.chat_input("Ask about ARCOS purchasing activity or DEA investigative history... "):
    ts_now = datetime.now().strftime("%H:%M")

    # Add user message
    st.session_state.messages.append({
        "role": "user", "content": prompt, "ts": ts_now
    })

    # Get response
    with st.spinner("Routing and querying..."):
        result = handle_query(prompt)

    answer = result.get("answer", "")
    intent = result.get("intent", "")
    sql    = result.get("sql")
    sql_explanation = result.get("sql_explanation")
    columns = result.get("columns") or []
    rows    = result.get("rows") or []

    # Add assistant message
    st.session_state.messages.append({
        "id":      str(uuid.uuid4()),
        "role":    "assistant",
        "content": answer,
        "intent":  intent,
        "sql":     sql,
        "sql_explanation": sql_explanation,
        "columns": columns,
        "rows":    rows,
        "ts":      ts_now,
    })

    st.rerun()